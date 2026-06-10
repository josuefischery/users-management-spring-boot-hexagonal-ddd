from __future__ import annotations

from datetime import date
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "report-assets"
OUT_DOCX = DOCS_DIR / "Informe_Actividad_Unidad_4_Cliente_RESTful.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
BORDER = "D9E2EC"
MUTED = "64748B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_cell_text(cell, bold: bool = False, color: str = "000000", size: int = 9) -> None:
    for paragraph in cell.paragraphs:
      paragraph.paragraph_format.space_after = Pt(0)
      for run in paragraph.runs:
          run.font.name = "Calibri"
          run.font.size = Pt(size)
          run.font.bold = bold
          run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_border(cell)
        style_cell_text(cell, bold=True, color=DARK_BLUE, size=9)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = value
            set_cell_width(cell, widths[index])
            set_cell_border(cell)
            style_cell_text(cell, size=9)

    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level <= 2 else DARK_BLUE)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.10


def add_code_image(document: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(6.2))
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(10)
        for run in caption_paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def generate_code_image(source: Path, output: Path, start: int, end: int, title: str) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()[start - 1 : end]
    numbered = [f"{idx + start:>3}  {line}" for idx, line in enumerate(lines)]

    try:
        font = ImageFont.truetype("/System/Library/Fonts/Menlo.ttc", 18)
        title_font = ImageFont.truetype("/System/Library/Fonts/Menlo.ttc", 20)
    except OSError:
        font = ImageFont.load_default()
        title_font = ImageFont.load_default()

    line_height = 26
    padding = 24
    header_height = 52
    width = 1560
    height = header_height + padding + line_height * len(numbered) + padding
    image = Image.new("RGB", (width, height), "#0f172a")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((0, 0, width - 1, height - 1), radius=18, outline="#334155", width=2)
    draw.rectangle((0, 0, width, header_height), fill="#111827")
    draw.text((padding, 16), title, font=title_font, fill="#e5e7eb")

    y = header_height + 18
    for line in numbered:
        draw.text((padding, y), line[:145], font=font, fill="#dbeafe")
        y += line_height

    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def build_document() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    code_assets = [
        (
            ROOT / "src/services/usersService.ts",
            ASSETS_DIR / "code-users-service.png",
            1,
            80,
            "Servicio HTTP usersService.ts",
        ),
        (
            ROOT / "src/utils/userValidation.ts",
            ASSETS_DIR / "code-validation.png",
            1,
            62,
            "Validaciones userValidation.ts",
        ),
        (
            ROOT / "src/models/user.ts",
            ASSETS_DIR / "code-model.png",
            1,
            34,
            "Modelo de datos user.ts",
        ),
        (
            ROOT / "src/components/UserForm.tsx",
            ASSETS_DIR / "code-form.png",
            1,
            110,
            "Formulario UserForm.tsx",
        ),
    ]

    for source, output, start, end, title in code_assets:
        generate_code_image(source, output, start, end, title)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Actividad Unidad 4\nCliente RESTful Web")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Desarrollo de Web - Ingenieria del software")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(24)

    metadata = [
        ("IES", "Universidad de Cartagena"),
        ("Tutor", "John Carlos Arrieta Arrieta"),
        ("Estudiante", "Josue David Fischer Yepes"),
        ("Modalidad", "Individual"),
        ("Framework", "React + TypeScript + Vite"),
        ("Fecha", date(2026, 6, 10).strftime("%d/%m/%Y")),
    ]
    add_table(document, ["Campo", "Descripcion"], metadata, [2200, 7160])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Repositorio backend base: ")
    r.font.bold = True
    p.add_run("https://github.com/arrietajohn/users-management-spring-boot-hexagonal-ddd")

    document.add_section(WD_SECTION.NEW_PAGE)

    add_heading(document, "Introduccion", 1)
    add_body(
        document,
        "El presente informe documenta el desarrollo de una aplicacion web frontend que funciona "
        "como cliente RESTful para el backend de gestion de usuarios construido con Spring Boot, "
        "arquitectura hexagonal y enfoque DDD. La solucion permite listar, consultar, crear, "
        "actualizar y eliminar registros desde una interfaz grafica, incorporando validaciones "
        "en el formulario, manejo de respuestas HTTP y mensajes de exito o error.",
    )

    add_heading(document, "Objetivos", 1)
    add_bullets(
        document,
        [
            "Construir un cliente RESTful web para consumir todos los endpoints expuestos por el backend.",
            "Implementar operaciones CRUD completas desde una interfaz grafica clara.",
            "Separar el codigo en componentes, servicios, modelos, utilidades de validacion y configuracion.",
            "Manejar respuestas del servidor, errores HTTP y validacion previa al envio de formularios.",
            "Documentar la integracion, estructura del proyecto y evidencias principales de la ejecucion.",
        ],
    )

    add_heading(document, "Framework seleccionado y justificacion", 1)
    add_body(
        document,
        "Se selecciono React con TypeScript y Vite. React permite organizar la interfaz en componentes "
        "reutilizables; TypeScript agrega tipado para los modelos y payloads del API; y Vite facilita "
        "un entorno de desarrollo rapido con proxy local hacia el backend. Esta combinacion cumple con "
        "los frameworks permitidos por la actividad y permite demostrar el consumo REST de forma clara.",
    )

    add_heading(document, "Descripcion general del backend consumido", 1)
    add_body(
        document,
        "El backend base es una aplicacion Spring Boot organizada con arquitectura hexagonal. La capa "
        "REST expone el recurso de usuarios bajo la ruta /api/users. Internamente el proyecto define "
        "casos de uso para crear, listar, consultar, actualizar y eliminar usuarios, y persiste la "
        "informacion en MySQL mediante la base de datos crud_usuarios.",
    )
    add_bullets(
        document,
        [
            "Puerto del backend: http://localhost:8080.",
            "URL base usada por el frontend: /api, redirigida por Vite hacia http://localhost:8080.",
            "Recurso principal: usuarios.",
            "Roles validos: ADMIN, MEMBER, REVIEWER.",
            "Estados validos: ACTIVE, INACTIVE, PENDING, BLOCKED.",
        ],
    )

    add_heading(document, "Tabla de endpoints consumidos", 1)
    endpoint_rows = [
        ["Listar usuarios", "GET", "/api/users", "Sin cuerpo", "200 OK, arreglo de usuarios"],
        ["Consultar usuario", "GET", "/api/users/{id}", "ID en la ruta", "200 OK usuario / 404 error"],
        [
            "Crear usuario",
            "POST",
            "/api/users",
            "id, name, email, password, role",
            "201 Created, usuario creado",
        ],
        [
            "Actualizar usuario",
            "PUT",
            "/api/users/{id}",
            "name, email, password, role, status",
            "200 OK, usuario actualizado",
        ],
        ["Eliminar usuario", "DELETE", "/api/users/{id}", "ID en la ruta", "204 No Content"],
    ]
    add_table(
        document,
        ["Operacion", "Metodo", "Endpoint", "Datos enviados", "Respuesta esperada"],
        endpoint_rows,
        [1900, 1100, 2100, 2500, 1760],
    )

    add_heading(document, "Estructura del proyecto frontend", 1)
    add_body(
        document,
        "La aplicacion fue organizada siguiendo buenas practicas de React: la vista principal orquesta "
        "el estado de la pantalla, los componentes se encargan de la presentacion, el servicio HTTP "
        "centraliza el consumo del API y las utilidades encapsulan validaciones y transformaciones.",
    )
    structure_rows = [
        ["src/models/user.ts", "Interfaces User, payloads de creacion/actualizacion, roles y estados."],
        ["src/services/usersService.ts", "Servicio REST con fetch, manejo de errores y metodos CRUD."],
        ["src/components/UserForm.tsx", "Formulario reutilizable para crear y editar usuarios."],
        ["src/components/UsersTable.tsx", "Tabla de registros con acciones de editar y eliminar."],
        ["src/components/ConfirmDialog.tsx", "Confirmacion antes de eliminar registros."],
        ["src/components/StatusMessage.tsx", "Mensajes de exito, informacion y error."],
        ["src/utils/userValidation.ts", "Validaciones de ID, nombre, correo, contrasena, rol y estado."],
        ["vite.config.ts", "Proxy /api hacia http://localhost:8080 para integracion local."],
    ]
    add_table(document, ["Archivo", "Responsabilidad"], structure_rows, [3000, 6360])

    add_heading(document, "Servicios REST implementados", 1)
    add_body(
        document,
        "El archivo usersService.ts define un cliente HTTP reutilizable. Cada metodo corresponde a una "
        "operacion del API: getAll, getById, create, update y remove. La funcion request centraliza "
        "cabeceras JSON, parseo de respuestas, manejo del codigo 204 y conversion de errores del "
        "backend al tipo ApiClientError para mostrarlos en la interfaz.",
    )
    add_code_image(document, ASSETS_DIR / "code-users-service.png", "Figura 1. Servicio HTTP encargado de consumir el API REST.")

    add_heading(document, "Modelos y validaciones", 1)
    add_body(
        document,
        "Los modelos TypeScript reflejan los campos usados por el backend: id, name, email, role y "
        "status. La validacion del formulario replica las restricciones principales del dominio: "
        "nombre minimo de tres caracteres, correo valido, contrasena minima de ocho caracteres y "
        "campos obligatorios antes de enviar peticiones al servidor.",
    )
    add_code_image(document, ASSETS_DIR / "code-model.png", "Figura 2. Interfaces y constantes usadas por el frontend.")
    add_code_image(document, ASSETS_DIR / "code-validation.png", "Figura 3. Validaciones del formulario antes del envio.")

    add_heading(document, "Formulario e interfaz grafica", 1)
    add_body(
        document,
        "La interfaz permite navegar entre las funcionalidades de listar, consultar por ID, crear, "
        "actualizar y eliminar usuarios. El formulario cambia entre modo creacion y edicion, bloquea "
        "el ID durante la edicion, muestra errores de validacion y permite seleccionar roles y estados "
        "validos para el backend.",
    )
    add_code_image(document, ASSETS_DIR / "code-form.png", "Figura 4. Componente de formulario para crear y editar usuarios.")

    app_shot = ASSETS_DIR / "app-main.png"
    add_heading(document, "Captura de ejecucion de la aplicacion", 1)
    if app_shot.exists():
        document.add_picture(str(app_shot), width=Inches(6.3))
        caption = document.add_paragraph("Figura 5. Interfaz principal del cliente RESTful en ejecucion local.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        for run in caption.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(MUTED)
    add_body(
        document,
        "La captura evidencia la pantalla principal con acciones de nuevo registro, actualizacion del "
        "listado, consulta por ID, tabla de usuarios, formulario y resumen de endpoints consumidos.",
    )

    add_heading(document, "Evidencia de operaciones CRUD", 1)
    add_table(
        document,
        ["Operacion", "Accion en la interfaz", "Endpoint usado", "Resultado esperado"],
        [
            ["Crear", "Diligenciar formulario y pulsar Crear usuario", "POST /api/users", "Mensaje de exito y registro en la tabla"],
            ["Consultar", "Ingresar ID y pulsar Buscar", "GET /api/users/{id}", "Detalle del usuario consultado"],
            ["Actualizar", "Pulsar editar, modificar campos y guardar", "PUT /api/users/{id}", "Datos actualizados en tabla y detalle"],
            ["Eliminar", "Pulsar eliminar y confirmar", "DELETE /api/users/{id}", "Registro eliminado y mensaje de exito"],
        ],
        [1600, 3200, 2300, 2260],
    )
    add_body(
        document,
        "Para registrar la evidencia final de la sustentacion se debe ejecutar el backend con MySQL y "
        "SMTP configurados, realizar las cuatro operaciones anteriores y capturar cada resultado. El "
        "frontend ya incluye las rutas, payloads y controles necesarios para ejecutar el ciclo CRUD.",
    )

    add_heading(document, "Problemas encontrados y solucion aplicada", 1)
    add_bullets(
        document,
        [
            "El backend no expone configuracion CORS visible en el repositorio. Se configuro proxy de Vite para consumir /api desde el mismo origen durante desarrollo.",
            "Cuando el backend no esta disponible, el proxy retorna 502. El servicio transforma ese estado en un mensaje claro para el usuario.",
            "Los formularios pueden enviar datos invalidos si no se validan en cliente. Se implemento validacion previa para nombre, correo, contrasena, rol y estado.",
            "La eliminacion es una accion irreversible. Se agrego un dialogo de confirmacion antes de ejecutar DELETE.",
        ],
    )

    add_heading(document, "Instrucciones de instalacion y ejecucion", 1)
    add_table(
        document,
        ["Paso", "Comando / accion"],
        [
            ["1", "Ejecutar el backend Spring Boot en http://localhost:8080."],
            ["2", "Configurar MySQL con la base crud_usuarios usando schema.sql del backend."],
            ["3", "En el frontend ejecutar npm install."],
            ["4", "Ejecutar npm run dev."],
            ["5", "Abrir http://localhost:5173 o la URL indicada por Vite."],
        ],
        [900, 8460],
    )

    add_heading(document, "Conclusiones", 1)
    add_bullets(
        document,
        [
            "La solucion implementa un cliente RESTful funcional y organizado para el recurso de usuarios del backend modelo.",
            "La separacion entre componentes, servicio HTTP, modelos y utilidades facilita el mantenimiento y cumple con la organizacion solicitada.",
            "El uso de TypeScript reduce errores al definir payloads y respuestas alineadas con el API.",
            "La interfaz permite demostrar las operaciones CRUD y maneja validaciones, confirmaciones y errores del servidor.",
        ],
    )

    add_heading(document, "Referencias consultadas", 1)
    add_bullets(
        document,
        [
            "Repositorio backend base: https://github.com/arrietajohn/users-management-spring-boot-hexagonal-ddd",
            "Documentacion de React: https://react.dev/",
            "Documentacion de Vite: https://vite.dev/",
            "Documentacion de Spring Boot: https://spring.io/projects/spring-boot",
            "Consigna de la Actividad Unidad 4 - Desarrollo de Web.",
        ],
    )

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Actividad Unidad 4 - Cliente RESTful Web")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(OUT_DOCX)
